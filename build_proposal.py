from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- Styles ---
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    if level == 1:
        h.font.size = Pt(22)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(16)
        h.font.bold = True
    else:
        h.font.size = Pt(13)
        h.font.bold = True

NAVY = RGBColor(0x1a, 0x3c, 0x6e)
TEAL = RGBColor(0x0d, 0x94, 0x8b)
DARK = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xff, 0xff, 0xff)

def add_colored_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A3C6E"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, text in enumerate(row_data):
            cell = row.cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            if r % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F7FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph('')
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AnfieldVoice')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Intelligent Visitor & Communication Management\nfor Residential Estates')
run.font.size = Pt(18)
run.font.color.rgb = TEAL
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for word, color in [('SECURE', NAVY), ('  \u25c6  ', RGBColor(0xcc, 0xcc, 0xcc)),
                     ('EFFICIENT', TEAL), ('  \u25c6  ', RGBColor(0xcc, 0xcc, 0xcc)),
                     ('BUILT FOR ESTATES', NAVY)]:
    run = p.add_run(word)
    run.font.size = Pt(14)
    run.font.color.rgb = color
    run.bold = True

for _ in range(8):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Confidential Proposal')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared July 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('')

toc_entries = [
    ('1.', 'Executive Summary'),
    ('2.', 'Company Profile'),
    ('3.', 'The Problem'),
    ('4.', 'Solution Overview'),
    ('5.', 'Product Features'),
    ('6.', 'Security Architecture'),
    ('7.', 'Visitor Management Workflows'),
    ('8.', 'Recurring Visitor Profiles'),
    ('9.', 'Delivery Management'),
    ('10.', 'Resident Communication'),
    ('11.', 'Administration Portal'),
    ('12.', 'Security Dashboard'),
    ('13.', 'System Architecture'),
    ('14.', 'Benefits & ROI'),
    ('15.', 'Competitive Comparison'),
    ('16.', 'Implementation Plan'),
    ('17.', 'Support & Maintenance'),
    ('18.', 'Frequently Asked Questions'),
    ('19.', 'Future Roadmap'),
    ('20.', 'Commercial Proposal'),
    ('21.', 'Conclusion'),
]

for num, title in toc_entries:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'AnfieldVoice is a purpose-built visitor management and resident communication platform '
    'designed exclusively for residential estates, body corporates, and gated communities. '
    'Unlike generic intercom replacements or VoIP systems, AnfieldVoice integrates visitor '
    'ticketing, recurring visitor profiles, delivery management, secure resident communication, '
    'and a dedicated security dashboard into a single, unified platform.'
)

doc.add_paragraph('The platform was designed around three principles:')

add_bullet(doc, 'Secure \u2014 layered visitor verification, complete audit trails, and privacy-first communication')
add_bullet(doc, 'Efficient \u2014 daily PIN authorisation, recurring visitor profiles, and reduced gate congestion')
add_bullet(doc, 'Built for Estates \u2014 apartment-focused workflows, body corporate administration, and integration-ready architecture')

doc.add_paragraph(
    'For a 360-unit residential estate, AnfieldVoice reduces gate processing time by an estimated '
    '60\u201380% for recurring visitors, eliminates approximately 40% of unnecessary resident-security '
    'phone calls, and provides a complete, searchable audit trail of every visitor, delivery, and '
    'access event. The platform pays for itself through operational efficiency alone, before '
    'accounting for the measurable improvement in resident satisfaction and security posture.'
)

doc.add_paragraph(
    'This proposal outlines the complete AnfieldVoice platform, including product features, '
    'security architecture, visitor management workflows, implementation methodology, and '
    'commercial terms for deployment at estates of 100 to 500+ residential units.'
)

doc.add_page_break()

# ============================================================
# 2. COMPANY PROFILE
# ============================================================
doc.add_heading('2. Company Profile', level=1)

doc.add_heading('About AnfieldVoice', level=2)
doc.add_paragraph(
    'AnfieldVoice was founded to solve a specific problem: residential estates need more than '
    'an intercom. They need a comprehensive visitor management and communication system that '
    'reflects the operational realities of modern estate living \u2014 daily domestic workers, school '
    'transport, delivery services, contractors, and guests, all moving through a single gate.'
)

doc.add_heading('Our Approach', level=2)
doc.add_paragraph(
    'We do not resell generic VoIP hardware with a custom label. AnfieldVoice is a vertically '
    'integrated solution \u2014 the mobile applications, server infrastructure, security dashboard, '
    'and administration portal are all designed and built as one system. The underlying '
    'telephony stack (SIP-based voice and video) uses battle-tested open-source components '
    '(FreeSWITCH, Kamailio) deployed on hardened, managed infrastructure.'
)

doc.add_heading('Technology Stack', level=2)

add_colored_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Mobile Apps', 'Custom-branded iOS & Android (Linphone core)', 'Resident & security communication'],
        ['SIP Proxy', 'Kamailio', 'Registration, routing, load balancing'],
        ['Media Server', 'FreeSWITCH', 'Voice & video call processing'],
        ['Database', 'PostgreSQL', 'Visitor records, audit trails, resident data'],
        ['API Layer', 'REST / WebSocket', 'Mobile provisioning, dashboards, integrations'],
        ['Visitor Terminal', 'Custom web application', 'Gate touchscreen for visitor check-in'],
        ['Security Dashboard', 'Real-time web interface', 'Gate officer view of all active visitors'],
    ]
)

doc.add_page_break()

# ============================================================
# 3. THE PROBLEM
# ============================================================
doc.add_heading('3. The Problem', level=1)

doc.add_heading('What Residential Estates Experience Today', level=2)

problems = [
    ('Gate Congestion',
     'Every morning, the same school transport, domestic workers, and delivery vehicles queue '
     'at the gate. Security processes each one as if they have never visited before. Residents '
     'receive repeated phone calls to verify the same people they approved yesterday and the day before.'),
    ('No Audit Trail',
     'When a resident disputes a delivery or questions who entered the estate, there is no '
     'centralised record. Gate logs, if they exist, are handwritten or stored in disconnected systems.'),
    ('Resident Frustration',
     'Residents are interrupted throughout the day with verification calls. Many stop answering, '
     'creating security gaps. Others answer and approve without verifying, defeating the purpose.'),
    ('Security Inefficiency',
     'Security officers spend 60\u201370% of their time on the phone with residents rather than '
     'monitoring the gate, patrolling the estate, or handling actual security incidents.'),
    ('Delivery Chaos',
     'Multiple delivery services arrive daily \u2014 couriers, grocery delivery, online shopping. '
     'Some estates ban deliveries at the gate; others accept them and create liability. '
     'Neither approach is satisfactory.'),
    ('No Resident Directory',
     'Residents have no easy way to contact neighbours, body corporate members, or estate '
     'management. Communication relies on WhatsApp groups, noticeboards, and word of mouth.'),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(desc)

doc.add_heading('The Cost of the Status Quo', level=2)
doc.add_paragraph(
    'For a 360-unit estate processing an estimated 200\u2013400 visitor arrivals per day, the '
    'cumulative cost of inefficient visitor management is substantial:'
)
add_bullet(doc, 'Security staff hours lost to phone calls: ~15\u201325 hours per week')
add_bullet(doc, 'Resident time lost to verification: ~2\u20135 minutes per call, multiple calls per day')
add_bullet(doc, 'Gate congestion during peak hours: 7:00\u20139:00 AM and 4:00\u20136:00 PM')
add_bullet(doc, 'Liability from undocumented deliveries and access events')
add_bullet(doc, 'Resident dissatisfaction with estate management \u2014 a top-3 complaint in body corporate surveys')

doc.add_page_break()

# ============================================================
# 4. SOLUTION OVERVIEW
# ============================================================
doc.add_heading('4. Solution Overview', level=1)

doc.add_paragraph(
    'AnfieldVoice replaces the estate intercom, visitor logbook, delivery register, and resident '
    'phone directory with a single, integrated platform. Every arrival, every call, and every '
    'access decision is recorded, searchable, and attributable.'
)

doc.add_heading('How It Works', level=2)

steps = [
    ('1. Resident Authorises a Visit',
     'Using the AnfieldVoice mobile app, a resident creates a Visitor Ticket: they enter '
     'the visitor\'s name, select a category, and optionally add vehicle details. The system '
     'generates a one-time PIN valid for today\'s date only. The resident shares this PIN '
     'with their visitor via SMS, WhatsApp, or any messaging app \u2014 directly from within AnfieldVoice.'),
    ('2. Visitor Arrives at the Gate',
     'The visitor approaches the Visitor Terminal at the gate \u2014 a weatherproof touchscreen '
     'kiosk or tablet. They enter their PIN. The system instantly retrieves their ticket '
     'and displays it on the Security Dashboard.'),
    ('3. Security Verifies and Admits',
     'The security officer sees the Visitor Ticket on their dashboard: who is expected, '
     'which resident authorised them, vehicle details if provided, and the ticket\'s validity. '
     'If everything matches, admission takes seconds. If something is different, the officer '
     'can call the resident directly from the dashboard with one click.'),
    ('4. Audit Trail Created',
     'Every step is logged: ticket creation, PIN validation, admission decision, security '
     'officer identity, and any exceptions. This data is available to body corporate '
     'administrators and never expires.'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = TEAL
    p.add_run(desc)

doc.add_heading('Platform Components', level=2)

add_colored_table(doc,
    ['Component', 'Users', 'Function'],
    [
        ['Resident Mobile App', 'Residents', 'Create tickets, call security/neighbours, manage recurring visitors, receive delivery notifications'],
        ['Visitor Terminal', 'Visitors at gate', 'Enter PIN, view instructions, call resident if needed'],
        ['Security Dashboard', 'Gate security officers', 'View active tickets, verify visitors, admit/deny, call residents, review audit log'],
        ['Administration Portal', 'Body corporate / estate management', 'Manage residents, configure access rules, view analytics, export audit reports'],
        ['AnfieldVoice Server', 'IT / managed service', 'SIP registration, call routing, ticket database, API, push notifications'],
    ]
)

doc.add_page_break()

# ============================================================
# 5. PRODUCT FEATURES
# ============================================================
doc.add_heading('5. Product Features', level=1)

doc.add_heading('5.1 Resident Mobile Application', level=2)
doc.add_paragraph(
    'A custom-branded application available on iOS and Android, built on the proven Linphone '
    'open-source telephony stack. The app provides:'
)
features_resident = [
    'Visitor Ticketing \u2014 Create one-time or recurring visitor tickets with PIN generation',
    'Recurring Visitor Profiles \u2014 Save trusted visitors for rapid daily authorisation',
    'Delivery Management \u2014 Receive notifications, approve or redirect deliveries',
    'Voice & Video Calling \u2014 Call security, neighbours, or estate management via SIP',
    'Resident Directory \u2014 Searchable directory of consenting residents',
    'Gate Intercom \u2014 Answer gate calls with video (optional camera integration)',
    'Push Notifications \u2014 Alerts for visitor arrivals, deliveries, and estate announcements',
    'Call History \u2014 Complete log of all communication within the estate',
    'Emergency Contact \u2014 One-touch call to security or emergency services',
]
for f in features_resident:
    add_bullet(doc, f)

doc.add_heading('5.2 Visitor Terminal', level=2)
doc.add_paragraph(
    'A weatherproof touchscreen interface deployed at the estate entrance gate(s):'
)
features_terminal = [
    'PIN Entry \u2014 Numeric keypad for visitor ticket PIN input',
    'Resident Search \u2014 Search by apartment number or resident name',
    'Direct Call \u2014 Initiate a SIP call to the resident\'s mobile app',
    'Visitor Instructions \u2014 Customisable welcome screen with estate rules',
    'QR Code Support \u2014 Scan QR codes from printed visitor passes (future)',
    'Ruggedised Hardware \u2014 IP65-rated enclosure for outdoor deployment',
]
for f in features_terminal:
    add_bullet(doc, f)

doc.add_heading('5.3 Security Dashboard', level=2)
doc.add_paragraph(
    'A real-time web interface for gate security officers:'
)
features_security = [
    'Live Ticket Queue \u2014 All active and upcoming visitor tickets for today',
    'Ticket Details \u2014 Resident name, apartment, visitor name, vehicle, category, PIN status',
    'Admit / Deny Controls \u2014 One-click admission logging with timestamp',
    'Recurring Visitor Flags \u2014 Visual indicator when a visitor has a verified profile',
    'One-Click Resident Call \u2014 Initiate a voice call to the resident directly from the ticket',
    'Audit Log \u2014 Searchable, filterable history of all access events',
    'Alert Notifications \u2014 Visual and audible alerts for expired tickets, denied access, or exceptions',
    'Multi-Language Support \u2014 Interface available in English, Afrikaans, Zulu, and other SA languages',
]
for f in features_security:
    add_bullet(doc, f)

doc.add_heading('5.4 Administration Portal', level=2)
doc.add_paragraph(
    'A web-based management console for body corporate administrators and estate managers:'
)
features_admin = [
    'Resident Management \u2014 Add, remove, and manage resident accounts and apartment assignments',
    'Visitor Policy Configuration \u2014 Set estate-wide rules for visitor types, time restrictions, and PIN expiry',
    'Analytics Dashboard \u2014 Visitor volume trends, peak hours, recurring visitor statistics, gate processing times',
    'Audit Export \u2014 Generate PDF/CSV reports of all access events for body corporate meetings or security reviews',
    'Billing & Usage \u2014 Track per-unit usage if operating on a cost-recovery model',
    'Communication Broadcasts \u2014 Send estate-wide announcements to all residents',
    'Role-Based Access \u2014 Granular permissions for administrators, security supervisors, and committee members',
]
for f in features_admin:
    add_bullet(doc, f)

doc.add_page_break()

# ============================================================
# 6. SECURITY ARCHITECTURE
# ============================================================
doc.add_heading('6. Security Architecture', level=1)

doc.add_paragraph(
    'AnfieldVoice is designed on a defence-in-depth security model. No single component '
    'failure or compromise can grant unauthorised access to the estate or resident data.'
)

doc.add_heading('6.1 Communication Security', level=2)
add_colored_table(doc,
    ['Layer', 'Protocol', 'Details'],
    [
        ['SIP Signalling', 'TLS 1.3', 'All SIP registration and call signalling encrypted in transit'],
        ['Media Streams', 'SRTP (AES-256)', 'Voice and video media encrypted end-to-end'],
        ['API Communication', 'HTTPS (TLS 1.3)', 'All REST and WebSocket traffic encrypted'],
        ['Push Notifications', 'APNs / FCM (encrypted)', 'Apple and Google push notification channels'],
        ['Database Connections', 'TLS + certificate pinning', 'All database queries over encrypted connections'],
    ]
)

doc.add_heading('6.2 Visitor Verification Layers', level=2)
doc.add_paragraph(
    'AnfieldVoice implements a multi-layered verification model. No single factor grants access:'
)

verification = [
    ('Layer 1 \u2014 Resident Authorisation',
     'Every visit begins with an explicit resident action: creating a ticket and generating a PIN. '
     'No visitor can be processed without prior resident authorisation.'),
    ('Layer 2 \u2014 Time-Bound PIN',
     'PINs are valid for today\'s date only. Expired PINs are rejected automatically. '
     'The resident sets the validity window (default: 24 hours from creation).'),
    ('Layer 3 \u2014 Ticket Retrieval',
     'When a PIN is entered at the gate, the system retrieves exactly one ticket. '
     'If the ticket does not exist or has been revoked, access is denied.'),
    ('Layer 4 \u2014 Security Officer Verification',
     'The security officer reviews the ticket details \u2014 visitor name, vehicle, category \u2014 '
     'and visually confirms the visitor before admission. The officer\'s identity is recorded.'),
    ('Layer 5 \u2014 Exception Handling',
     'If any detail does not match (wrong vehicle, unexpected time, different person), the '
     'officer can call the resident directly from the dashboard. No admission without resolution.'),
    ('Layer 6 \u2014 Audit Trail',
     'Every decision is logged immutably. The audit trail includes: who created the ticket, '
     'when the PIN was used, which officer processed the visitor, the admission decision, '
     'and any exceptions.'),
]

for title, desc in verification:
    p = doc.add_paragraph()
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(desc)

doc.add_heading('6.3 Data Security', level=2)
add_bullet(doc, 'Resident data is encrypted at rest (AES-256) and in transit (TLS 1.3)')
add_bullet(doc, 'PINs are hashed (bcrypt) \u2014 never stored in plaintext')
add_bullet(doc, 'Database access is restricted by IP whitelist and certificate authentication')
add_bullet(doc, 'All administrative access is logged and auditable')
add_bullet(doc, 'POPIA-compliant data handling with configurable retention policies')
add_bullet(doc, 'Regular automated security updates for all server components')

doc.add_page_break()

# ============================================================
# 7. VISITOR MANAGEMENT WORKFLOWS
# ============================================================
doc.add_heading('7. Visitor Management Workflows', level=1)

doc.add_heading('7.1 Standard Visitor (One-Time)', level=2)
doc.add_paragraph(
    'For a dinner guest, family member, or service provider visiting once:'
)

std_workflow = [
    'Resident opens AnfieldVoice app \u2192 "New Visitor Ticket"',
    'Enters visitor name, selects category (Guest, Contractor, Service, etc.)',
    'Optionally adds vehicle details for faster gate processing',
    'System generates today\'s PIN (e.g., 483921)',
    'Resident shares PIN with visitor via the app (SMS, WhatsApp, or copy)',
    'Visitor arrives at gate \u2192 enters PIN on Visitor Terminal',
    'Security sees ticket on dashboard \u2192 verifies \u2192 admits',
    'Ticket marked as "Admitted" with timestamp and officer ID',
    'System records complete audit trail entry',
]
for step in std_workflow:
    add_bullet(doc, step)

doc.add_heading('7.2 Recurring Visitor (Streamlined)', level=2)
doc.add_paragraph(
    'For a domestic worker, scholar transport, or gardener who visits daily:'
)

rec_workflow = [
    'Resident creates a Recurring Visitor Profile once (name, category, vehicle, schedule)',
    'Each morning, resident opens app \u2192 selects profile \u2192 "Generate Today\'s PIN"',
    'System generates PIN \u2014 resident shares it (or visitor knows to collect it daily)',
    'Visitor arrives at gate \u2192 enters PIN',
    'Security dashboard shows "Recurring Visitor \u2014 Verified Profile" with full details',
    'Officer visually confirms vehicle and person match the profile',
    'If match: quick admission. If mismatch: officer calls resident.',
    'Admission logged with recurring visitor reference',
]
for step in rec_workflow:
    add_bullet(doc, step)

doc.add_heading('7.3 Unexpected Visitor', level=2)
doc.add_paragraph(
    'For a visitor who arrives without a PIN:'
)

unexp_workflow = [
    'Visitor approaches gate \u2192 does not have a PIN',
    'Visitor uses terminal to search for resident by name or apartment number',
    'Terminal initiates a SIP call to the resident\'s mobile app',
    'Resident receives call with caller ID "Estate Gate"',
    'If resident answers and approves, they can generate a PIN on the spot via the app',
    'Visitor enters PIN \u2192 standard verification \u2192 admission',
    'If resident does not answer, visitor is not admitted',
    'Event is logged as "Unexpected \u2014 No PIN \u2014 No Resident Response"',
]
for step in unexp_workflow:
    add_bullet(doc, step)

doc.add_heading('7.4 Denied Access', level=2)
doc.add_paragraph('When a visitor is denied admission, the system records:')
add_bullet(doc, 'Date and time of attempt')
add_bullet(doc, 'PIN entered (if any) and validation result')
add_bullet(doc, 'Reason for denial (expired PIN, revoked ticket, resident declined, no response)')
add_bullet(doc, 'Security officer who processed the denial')
add_bullet(doc, 'Any resident communication attempt and outcome')

doc.add_page_break()

# ============================================================
# 8. RECURRING VISITOR PROFILES
# ============================================================
doc.add_heading('8. Recurring Visitor Profiles', level=1)

doc.add_paragraph(
    'This is one of AnfieldVoice\'s strongest differentiators. Most residential estates receive '
    'the same visitors every day \u2014 scholar transport, domestic workers, caregivers, garden services, '
    'contractors, and regular couriers. Traditional intercom systems treat every arrival as a new '
    'visitor, forcing security to repeat the same verification process daily.'
)

p = doc.add_paragraph()
run = p.add_run('AnfieldVoice introduces Recurring Visitor Profiles, ')
run.bold = True
p.add_run(
    'allowing residents to register trusted, regularly visiting individuals while still '
    'maintaining daily security controls.'
)

doc.add_heading('8.1 Profile Structure', level=2)

add_colored_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['Full Name', 'Required', 'Visitor\'s full name as it should appear on tickets'],
        ['Mobile Number', 'Optional', 'For automated PIN delivery (future feature)'],
        ['Category', 'Required', 'Scholar Transport, Staff Transport, Domestic Worker, Gardener, Caregiver, Contractor, Other'],
        ['Vehicle Make', 'Optional', 'e.g., Toyota, Volkswagen, Ford'],
        ['Vehicle Model', 'Optional', 'e.g., Quantum, Polo, Ranger'],
        ['Vehicle Colour', 'Optional', 'For visual verification at gate'],
        ['Registration Number', 'Optional', 'e.g., CA 123-456'],
        ['Typical Arrival Days', 'Optional', 'Monday\u2013Friday, Weekends, or specific days'],
        ['Typical Arrival Times', 'Optional', 'Morning (6\u20139 AM), Afternoon, Evening, or specific time window'],
        ['Notes', 'Optional', 'Free-text field for any additional information'],
    ]
)

doc.add_heading('8.2 Daily Workflow', level=2)
doc.add_paragraph(
    'Rather than completing the full visitor form every day, the resident simply:'
)
add_bullet(doc, 'Opens the AnfieldVoice app')
add_bullet(doc, 'Selects "Use Existing Profile"')
add_bullet(doc, 'Chooses the relevant recurring visitor from their saved list')
add_bullet(doc, 'Taps "Generate Today\'s PIN"')
doc.add_paragraph(
    'The entire process takes under 5 seconds. The PIN is still required \u2014 this confirms that '
    'the resident expects this person today, preventing unauthorised entry by someone who was '
    'admitted yesterday.'
)

doc.add_heading('8.3 Security Benefits', level=2)
add_bullet(doc, 'Vehicle and registration are already recorded \u2014 no verbal description at the gate')
add_bullet(doc, 'Resident has previously approved the visitor \u2014 profile creation is a deliberate act')
add_bullet(doc, 'Security becomes familiar with the individual over time \u2014 anomalies stand out')
add_bullet(doc, 'Daily PIN requirement prevents assumption-based access')
add_bullet(doc, 'Complete audit trail links every visit back to the resident\'s explicit authorisation')

doc.add_heading('8.4 Operational Benefits', level=2)
add_bullet(doc, 'Gate processing time for recurring visitors: ~10\u201315 seconds vs. 60\u201390 seconds for new visitors')
add_bullet(doc, 'Reduced resident interruption: no daily verification calls for known visitors')
add_bullet(doc, 'Morning peak congestion reduced by an estimated 40\u201360%')
add_bullet(doc, 'Security officers can focus on unfamiliar visitors and actual security duties')

p = doc.add_paragraph()
run = p.add_run('\nMarketing Position: ')
run.bold = True
run.font.color.rgb = TEAL
p.add_run(
    '"Designed for real estate operations \u2014 not just visitors." This differentiates AnfieldVoice '
    'from generic intercom replacements that treat every arrival identically.'
)

doc.add_page_break()

# ============================================================
# 9. DELIVERY MANAGEMENT
# ============================================================
doc.add_heading('9. Delivery Management', level=1)

doc.add_paragraph(
    'Delivery management is one of the most persistent challenges for residential estates. '
    'Multiple courier services, grocery deliveries, and online shopping orders arrive daily. '
    'AnfieldVoice provides a structured workflow that balances resident convenience with '
    'estate security.'
)

doc.add_heading('9.1 Delivery Workflow', level=2)

delivery_steps = [
    ('Notification',
     'Resident receives a push notification when a delivery is en route (integration with '
     'delivery service APIs or manual entry by security).'),
    ('Pre-Approval',
     'Resident can pre-approve the delivery from the app: "Accept," "Hold at Gate," or "Redirect." '
     'If no response, estate policy determines default handling.'),
    ('Arrival',
     'Delivery driver identifies the delivery at the Visitor Terminal. Security sees the '
     'resident\'s pre-approval status on the dashboard.'),
    ('Processing',
     'Approved: driver proceeds to the apartment or designated drop-off point. '
     'Hold at Gate: package is logged and stored at security. '
     'Redirect: resident specifies an alternative (leave with neighbour, deliver tomorrow, etc.).'),
    ('Confirmation',
     'Resident receives delivery confirmation with timestamp. Delivery event is logged in '
     'the audit trail with driver details, time, and resolution.'),
]

for title, desc in delivery_steps:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(desc)

doc.add_heading('9.2 Benefits', level=2)
add_bullet(doc, 'Reduces gate congestion by pre-processing delivery decisions')
add_bullet(doc, 'Eliminates "we have a package for apartment X" phone calls')
add_bullet(doc, 'Creates accountability \u2014 every delivery is logged with driver details')
add_bullet(doc, 'Resident maintains control without being interrupted')
add_bullet(doc, 'Estate policy is consistently applied, not dependent on which security officer is on duty')

doc.add_page_break()

# ============================================================
# 10. RESIDENT COMMUNICATION
# ============================================================
doc.add_heading('10. Resident Communication', level=1)

doc.add_paragraph(
    'AnfieldVoice is not just a visitor management system \u2014 it is the estate\'s communication '
    'backbone. It replaces the fragmented mix of intercoms, WhatsApp groups, noticeboards, '
    'and phone trees with a single, private network.'
)

doc.add_heading('10.1 Communication Features', level=2)

add_colored_table(doc,
    ['Feature', 'Description'],
    [
        ['Voice Calling', 'SIP-based HD voice calls between residents, security, and estate management. No cellular minutes used within the estate network.'],
        ['Video Calling', 'H.264/H.265 video calls between residents and the gate terminal. Optional camera integration at the gate.'],
        ['Resident Directory', 'Searchable directory of consenting residents by name or apartment number. Opt-in \u2014 residents control their visibility.'],
        ['Gate Intercom', 'Visitors without a PIN can call the resident directly from the Visitor Terminal. Caller ID displays "Estate Gate."'],
        ['Estate Broadcasts', 'Body corporate or estate management can send announcements to all residents or specific buildings.'],
        ['Emergency Calling', 'One-touch call to security with priority routing. Optionally configurable for armed response or medical emergency.'],
        ['Call History', 'Complete log of all calls with timestamps, accessible from the resident app.'],
    ]
)

doc.add_heading('10.2 Privacy Model', level=2)
doc.add_paragraph(
    'AnfieldVoice operates as a closed network. Residents are not reachable from outside '
    'the estate unless they explicitly enable external calling. By default:'
)
add_bullet(doc, 'Residents can call security, the gate, and other consenting residents')
add_bullet(doc, 'Security and the gate can call residents')
add_bullet(doc, 'Residents cannot be called from outside the estate network (configurable)')
add_bullet(doc, 'Resident phone numbers are never exposed to other residents \u2014 calls use apartment numbers or names')
add_bullet(doc, 'All communication is encrypted (TLS + SRTP)')

doc.add_page_break()

# ============================================================
# 11. ADMINISTRATION PORTAL
# ============================================================
doc.add_heading('11. Administration Portal', level=1)

doc.add_paragraph(
    'The Administration Portal is the central management console for body corporate '
    'administrators, estate managers, and security supervisors. It provides complete '
    'control over the AnfieldVoice deployment without requiring technical expertise.'
)

doc.add_heading('11.1 Resident Management', level=2)
add_bullet(doc, 'Add, edit, and deactivate resident accounts')
add_bullet(doc, 'Assign residents to apartments / units')
add_bullet(doc, 'Manage resident directory visibility preferences')
add_bullet(doc, 'Bulk import from existing resident databases (CSV, Excel)')
add_bullet(doc, 'Track account activation status and app installation')

doc.add_heading('11.2 Visitor Policy Configuration', level=2)
add_bullet(doc, 'Set PIN validity duration (default: 24 hours)')
add_bullet(doc, 'Configure visitor categories available to residents')
add_bullet(doc, 'Set maximum daily visitor tickets per unit')
add_bullet(doc, 'Define restricted hours (e.g., no visitors after 22:00)')
add_bullet(doc, 'Configure delivery handling policies')

doc.add_heading('11.3 Analytics & Reporting', level=2)
add_bullet(doc, 'Visitor volume dashboard \u2014 daily, weekly, monthly trends')
add_bullet(doc, 'Peak hour analysis for security staffing optimisation')
add_bullet(doc, 'Recurring visitor statistics by category')
add_bullet(doc, 'Gate processing time averages and outliers')
add_bullet(doc, 'Denied access events with reasons')
add_bullet(doc, 'Security officer performance metrics')
add_bullet(doc, 'Export reports as PDF or CSV for body corporate meetings')

doc.add_heading('11.4 Audit & Compliance', level=2)
add_bullet(doc, 'Complete, searchable audit trail of all access events')
add_bullet(doc, 'Filter by date range, resident, visitor, category, or outcome')
add_bullet(doc, 'Immutable logs \u2014 events cannot be modified or deleted')
add_bullet(doc, 'Configurable data retention policies (POPIA-compliant)')
add_bullet(doc, 'Audit export for security incidents or insurance claims')

doc.add_page_break()

# ============================================================
# 12. SECURITY DASHBOARD
# ============================================================
doc.add_heading('12. Security Dashboard', level=1)

doc.add_paragraph(
    'The Security Dashboard is the primary interface for gate security officers. It is '
    'designed for speed and clarity \u2014 an officer can process a visitor in seconds without '
    'training on complex software.'
)

doc.add_heading('12.1 Dashboard Layout', level=2)

doc.add_paragraph(
    'The dashboard presents a real-time, auto-refreshing view of all active and upcoming '
    'visitor tickets for the current day:'
)

add_colored_table(doc,
    ['Column', 'Content'],
    [
        ['Time', 'Scheduled or expected arrival window'],
        ['Resident', 'Resident name and apartment number'],
        ['Visitor', 'Visitor name and category label'],
        ['Vehicle', 'Vehicle details (if provided) \u2014 make, model, colour, registration'],
        ['Type', 'Standard or Recurring (with verified profile badge)'],
        ['Status', 'Pending, Arrived, Admitted, Denied, Expired'],
        ['PIN', 'Ticket PIN \u2014 validated on entry at the terminal'],
        ['Actions', 'Admit, Deny, Call Resident buttons'],
    ]
)

doc.add_heading('12.2 Key Capabilities', level=2)
add_bullet(doc, 'Real-time ticket queue with colour-coded status indicators')
add_bullet(doc, '"Recurring Visitor \u2014 Verified Profile" badge for pre-registered visitors')
add_bullet(doc, 'One-click resident calling from any ticket')
add_bullet(doc, 'Search: find tickets by resident name, visitor name, apartment, or PIN')
add_bullet(doc, 'Alert panel: expired tickets, denied access, exceptions requiring attention')
add_bullet(doc, 'Multi-language interface (English, Afrikaans, Zulu, and others)')
add_bullet(doc, 'Works on standard office monitors, tablets, or ruggedised touchscreens')
add_bullet(doc, 'Low-bandwidth operation \u2014 functions reliably on basic internet connections')

doc.add_heading('12.3 Officer Accountability', level=2)
add_bullet(doc, 'Each officer logs in with unique credentials')
add_bullet(doc, 'All admission/denial decisions are attributed to the logged-in officer')
add_bullet(doc, 'Shift handover: on-duty officer is tracked for audit purposes')
add_bullet(doc, 'Supervisor view: monitor multiple gate stations from a single interface')

doc.add_page_break()

# ============================================================
# 13. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('13. System Architecture', level=1)

doc.add_paragraph(
    'AnfieldVoice is built on a modular, open-standards architecture that ensures '
    'reliability, scalability, and future extensibility.'
)

doc.add_heading('13.1 Architecture Overview', level=2)

arch_text = (
    '                            RESIDENT MOBILE APPS\n'
    '        +-------------+  +-------------+  +-------------+\n'
    '        |   iOS App   |  | Android App |  | Desktop App |\n'
    '        | (Linphone)  |  | (Linphone)  |  | (Linphone)  |\n'
    '        +------+------+  +------+------+  +------+------+\n'
    '               |                |                |\n'
    '               +----------------+----------------+\n'
    '                                |  TLS 1.3\n'
    '                                |\n'
    '                    ANFIELDVOICE SERVER\n'
    '        +---------------------------------------------+\n'
    '        |          KAMAILIO (SIP Proxy)               |\n'
    '        |   Registration . Routing . Load Balance     |\n'
    '        +--------------------+------------------------+\n'
    '                             |\n'
    '        +--------------------+------------------------+\n'
    '        |          FREESWITCH (Media Server)          |\n'
    '        |   Voice . Video . Conferencing . IVR        |\n'
    '        +--------------------+------------------------+\n'
    '                             |\n'
    '        +--------------------+------------------------+\n'
    '        |          POSTGRESQL DATABASE                |\n'
    '        |   Residents . Tickets . Audit . Config      |\n'
    '        +--------------------+------------------------+\n'
    '                             |\n'
    '        +--------------------+------------------------+\n'
    '        |          REST API / WebSocket               |\n'
    '        |   Mobile Provisioning . Dashboards . Admin  |\n'
    '        +--------------------+------------------------+\n'
    '                             |\n'
    '                     ESTATE INFRASTRUCTURE\n'
    '        +-------------+  +-------------+  +-------------+\n'
    '        |   Visitor   |  |  Security   |  |    Admin    |\n'
    '        |  Terminal   |  |  Dashboard  |  |   Portal    |\n'
    '        |   (Gate)    |  | (Gate/Off.) |  | (Remote)    |\n'
    '        +-------------+  +-------------+  +-------------+'
)

p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(7.5)
run.font.color.rgb = DARK

doc.add_heading('13.2 Technology Stack Details', level=2)

add_colored_table(doc,
    ['Component', 'Technology', 'Rationale'],
    [
        ['SIP Proxy', 'Kamailio 5.x', 'Industry-standard SIP proxy. Handles 10,000+ concurrent registrations.'],
        ['Media Server', 'FreeSWITCH 1.10+', 'Handles voice/video transcoding, conferencing, IVR.'],
        ['Database', 'PostgreSQL 16+', 'ACID-compliant, mature, excellent replication and backup.'],
        ['API Server', 'FastAPI / Node.js', 'REST + WebSocket for apps, dashboards, integrations.'],
        ['Mobile Apps', 'Linphone SDK', 'Open-source SIP client. Customisable UI. iOS, Android, desktop.'],
        ['Visitor Terminal', 'PWA', 'Works on any modern browser. Deploy on ruggedised hardware.'],
        ['Dashboards', 'React / Vue.js SPA', 'Real-time web interfaces. Lightweight, low-spec friendly.'],
        ['Push Notifications', 'APNs + FCM', 'Apple and Google push for background call delivery.'],
    ]
)

doc.add_heading('13.3 Deployment Options', level=2)
add_bullet(doc, 'On-Premises \u2014 Server hardware deployed in the estate\'s network room or security office')
add_bullet(doc, 'Cloud-Hosted \u2014 AnfieldVoice-managed cloud deployment (AWS Johannesburg region)')
add_bullet(doc, 'Hybrid \u2014 Core server on-premises with cloud backup and failover')

doc.add_heading('13.4 Scalability', level=2)
doc.add_paragraph('The architecture is designed to scale well beyond 360 units:')
add_bullet(doc, 'Kamailio handles 10,000+ concurrent SIP registrations on modest hardware')
add_bullet(doc, 'FreeSWITCH scales horizontally \u2014 additional media servers for larger estates')
add_bullet(doc, 'PostgreSQL supports read replicas for reporting without impacting transactions')
add_bullet(doc, 'Same architecture serves estates from 50 to 5,000+ units without fundamental changes')

doc.add_page_break()

# ============================================================
# 14. BENEFITS & ROI
# ============================================================
doc.add_heading('14. Benefits & ROI', level=1)

doc.add_heading('14.1 Quantitative Benefits', level=2)

doc.add_paragraph(
    'Estimated impact for a 360-unit estate processing approximately 300 visitor/delivery '
    'arrivals per day:'
)

add_colored_table(doc,
    ['Metric', 'Before AnfieldVoice', 'With AnfieldVoice', 'Improvement'],
    [
        ['Gate processing (standard visitor)', '60\u201390 seconds', '20\u201330 seconds', '~60% reduction'],
        ['Gate processing (recurring visitor)', '60\u201390 seconds', '10\u201315 seconds', '~80% reduction'],
        ['Resident verification calls/day', '50\u201380 calls', '10\u201320 calls', '~70% reduction'],
        ['Security staff time on phone', '15\u201325 hrs/week', '3\u20135 hrs/week', '~75% reduction'],
        ['Morning peak queue length', '8\u201315 vehicles', '2\u20135 vehicles', '~60% reduction'],
        ['Missed/undocumented deliveries', '5\u201315 per month', '< 1 per month', '~95% reduction'],
        ['Audit trail availability', 'Manual / none', 'Complete, searchable', 'New capability'],
    ]
)

doc.add_heading('14.2 Qualitative Benefits', level=2)
add_bullet(doc, 'Improved resident satisfaction \u2014 faster gate processing, fewer interruptions')
add_bullet(doc, 'Enhanced security posture \u2014 layered verification, complete audit trail')
add_bullet(doc, 'Professional estate image \u2014 branded app, modern visitor terminal')
add_bullet(doc, 'Security officer job satisfaction \u2014 less phone work, more actual security')
add_bullet(doc, 'Body corporate transparency \u2014 data-driven reports for AGMs and committee meetings')
add_bullet(doc, 'Reduced liability \u2014 documented access decisions for insurance and dispute resolution')

doc.add_heading('14.3 Return on Investment', level=2)
doc.add_paragraph('AnfieldVoice delivers ROI through multiple channels:')
add_bullet(doc, 'Operational Efficiency \u2014 Security staff reallocation from phone duty to active security/patrol')
add_bullet(doc, 'Risk Reduction \u2014 Documented audit trail reduces liability exposure and insurance costs')
add_bullet(doc, 'Resident Retention \u2014 Improved living experience contributes to property value')
add_bullet(doc, 'Process Automation \u2014 Eliminates manual logbooks, paper forms, and handwritten registers')
doc.add_paragraph(
    'For a 360-unit estate, the operational savings alone typically recover the platform cost '
    'within 12\u201318 months, depending on deployment configuration and security staffing model.'
)

doc.add_page_break()

# ============================================================
# 15. COMPETITIVE COMPARISON
# ============================================================
doc.add_heading('15. Competitive Comparison', level=1)

doc.add_paragraph(
    'AnfieldVoice competes across two categories: traditional intercom/VoIP vendors and '
    'modern visitor management platforms.'
)

add_colored_table(doc,
    ['Capability', 'AnfieldVoice', 'ButterflyMX', '2N / Akuvox', 'Generic VoIP', 'Manual/Paper'],
    [
        ['Visitor Ticketing', '\u2713 Full', '\u2713 Basic', '\u2717', '\u2717', '\u2717'],
        ['PIN-Based Access', '\u2713', '\u2713', '\u2717', '\u2717', '\u2717'],
        ['Recurring Visitor Profiles', '\u2713 Native', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['Delivery Management', '\u2713 Full', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['Resident Mobile App', '\u2713 Branded', '\u2713', '\u2717', '\u2717', '\u2717'],
        ['Video Calling', '\u2713', '\u2713', '\u2713', '\u2713', '\u2717'],
        ['Security Dashboard', '\u2713 Purpose-built', '\u2713 Basic', '\u2717', '\u2717', '\u2717'],
        ['Complete Audit Trail', '\u2713', '\u2713 Limited', '\u2717', '\u2717', '\u2717'],
        ['Administration Portal', '\u2713 Full', '\u2713 Limited', '\u2713 Basic', '\u2717', '\u2717'],
        ['Resident Directory', '\u2713', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['Estate Broadcasts', '\u2713', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['Custom Branding', '\u2713 Full', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['Multi-Language (SA)', '\u2713', '\u2717', '\u2717', '\u2717', '\u2717'],
        ['On-Premises Option', '\u2713', '\u2717', '\u2713', '\u2713', 'N/A'],
        ['POPIA Compliance', '\u2713 Built-in', 'Partial', 'N/A', 'N/A', '\u2717'],
        ['SA-Based Support', '\u2713 Local', '\u2717 US-based', '\u2717 EU-based', 'Varies', 'N/A'],
        ['Open Standards', '\u2713 SIP/REST', 'Proprietary', 'Proprietary', '\u2713 SIP', 'N/A'],
        ['Scalability (500+)', '\u2713 Native', '\u2713', '\u2713', 'Partial', '\u2717'],
    ]
)

doc.add_heading('15.1 Key Differentiators', level=2)

diff = [
    ('Recurring Visitor Profiles',
     'No competitor offers native, purpose-built recurring visitor management. This feature '
     'alone transforms daily gate operations for estates with domestic workers, scholar '
     'transport, and regular service providers.'),
    ('Built for Estates, Not Adapted',
     'ButterflyMX and similar platforms were designed for US apartment buildings with '
     'different operational models. AnfieldVoice was designed for South African residential '
     'estates \u2014 gated communities with body corporate governance, security officers, and '
     'the specific visitor categories relevant to SA estate living.'),
    ('Complete Platform, Not a Component',
     'AnfieldVoice is not a VoIP system with visitor management added, nor a visitor '
     'management system with VoIP added. Both were designed together as one platform. '
     'The result is a seamless experience: the security dashboard, visitor terminal, '
     'resident app, and administration portal all share the same data model.'),
    ('Open Standards, Not Vendor Lock-In',
     'Built on SIP, TLS/SRTP, PostgreSQL, and REST APIs \u2014 industry-standard protocols '
     'and open-source components. The estate is never locked into proprietary hardware '
     'or protocols.'),
]

for title, desc in diff:
    p = doc.add_paragraph()
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = TEAL
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 16. IMPLEMENTATION PLAN
# ============================================================
doc.add_heading('16. Implementation Plan', level=1)

doc.add_paragraph(
    'AnfieldVoice follows a phased deployment methodology designed to minimise disruption '
    'to estate operations while ensuring a successful rollout.'
)

doc.add_heading('16.1 Phased Rollout', level=2)

add_colored_table(doc,
    ['Phase', 'Duration', 'Activities', 'Deliverables'],
    [
        ['1. Discovery & Planning', '2 weeks',
         'Site survey, network assessment, gate hardware evaluation, resident database collection, branding requirements',
         'Technical assessment report, deployment plan, branding mockups'],
        ['2. Infrastructure Setup', '2\u20133 weeks',
         'Server deployment, SIP trunk configuration, network integration, security hardening',
         'Operational server environment, network connectivity verified'],
        ['3. Platform Configuration', '2 weeks',
         'Branded app build, visitor terminal setup, dashboard configuration, resident data import',
         'Branded apps ready, configured dashboards, imported resident data'],
        ['4. Security Staff Training', '1 week',
         'Dashboard training, visitor processing workflows, exception handling, shift handover procedures',
         'Trained security staff, training documentation, quick-reference guides'],
        ['5. Resident Onboarding', '2\u20133 weeks',
         'Resident communication, app installation support, training sessions, FAQ distribution',
         '>80% resident app installation, resident training materials'],
        ['6. Go-Live & Parallel Run', '1\u20132 weeks',
         'Dual operation with existing system, monitoring, rapid response, policy adjustment',
         'Go-live completion report, issue resolution log'],
        ['7. Post-Go-Live Optimisation', '2 weeks',
         'Performance review, policy tuning, resident feedback collection, additional training',
         'Optimisation report, resident satisfaction survey results'],
    ]
)

doc.add_heading('16.2 Total Timeline', level=2)
doc.add_paragraph(
    'Estimated total: 10\u201314 weeks from contract signing to full operational deployment. '
    'Timeline varies based on estate size, existing infrastructure, and resident engagement.'
)

doc.add_heading('16.3 Resident Adoption Strategy', level=2)
doc.add_paragraph(
    'Resident adoption is the critical success factor. AnfieldVoice includes a structured '
    'adoption programme:'
)
add_bullet(doc, 'Pre-launch communication: body corporate notice, welcome letter, FAQ document')
add_bullet(doc, 'Launch event: on-site app installation support, live demonstration')
add_bullet(doc, 'App Store / Google Play: branded app listing with estate-specific instructions')
add_bullet(doc, 'Onboarding wizard: first-launch walkthrough, account activation, directory opt-in')
add_bullet(doc, 'Ongoing support: in-app help, resident support email/phone, quarterly tips newsletter')
add_bullet(doc, 'Incentive: early-adopter recognition, reduced gate wait times as immediate benefit')

doc.add_page_break()

# ============================================================
# 17. SUPPORT & MAINTENANCE
# ============================================================
doc.add_heading('17. Support & Maintenance', level=1)

doc.add_heading('17.1 Support Tiers', level=2)

add_colored_table(doc,
    ['Tier', 'Response Time', 'Hours', 'Channels', 'Coverage'],
    [
        ['Critical \u2014 System down', '< 1 hour', '24/7/365', 'Phone + Email + Chat', 'Server, SIP, database'],
        ['High \u2014 Major feature unavailable', '< 4 hours', 'Business hrs + Sat', 'Email + Ticketing', 'Dashboard, terminal, app'],
        ['Medium \u2014 Minor issue', '< 1 business day', 'Business hours', 'Email + Ticketing', 'Non-critical features'],
        ['Low \u2014 Question / Enhancement', '< 3 business days', 'Business hours', 'Email', 'General inquiries'],
    ]
)

doc.add_heading('17.2 Maintenance Services', level=2)
add_bullet(doc, 'Automated security updates for all server components (weekly)')
add_bullet(doc, 'Database backups (hourly incremental, daily full, encrypted off-site)')
add_bullet(doc, 'SIP infrastructure monitoring (registration health, call quality metrics)')
add_bullet(doc, 'Mobile app updates (iOS App Store and Google Play releases)')
add_bullet(doc, 'Quarterly platform health review with body corporate / estate management')
add_bullet(doc, 'Annual disaster recovery test')

doc.add_heading('17.3 Service Level Agreement', level=2)
add_bullet(doc, '99.5% platform uptime guarantee (excluding scheduled maintenance)')
add_bullet(doc, 'Scheduled maintenance: Sunday 02:00\u201304:00, with 72-hour advance notice')
add_bullet(doc, 'Monthly uptime report included in management dashboard')
add_bullet(doc, 'Penalty clauses for SLA breaches (negotiated per deployment agreement)')

doc.add_page_break()

# ============================================================
# 18. FAQ
# ============================================================
doc.add_heading('18. Frequently Asked Questions', level=1)

faqs = [
    ('What hardware does the estate need to provide?',
     'AnfieldVoice requires: (1) a visitor terminal at the gate \u2014 typically a weatherproof '
     'touchscreen tablet or all-in-one PC, (2) a monitor/computer for the security dashboard, '
     '(3) a server or cloud subscription for the AnfieldVoice platform. We provide hardware '
     'recommendations and procurement assistance. The estate\'s existing internet connection '
     'is sufficient in most cases.'),
    ('Does this replace our existing intercom system?',
     'Yes. AnfieldVoice replaces the intercom, visitor logbook, and resident directory with '
     'an integrated digital platform. If the estate has existing VoIP hardware (e.g., SIP '
     'door stations), we can often integrate with it.'),
    ('What happens if the internet goes down?',
     'The system includes an offline fallback mode. The visitor terminal caches today\'s '
     'tickets locally. If the internet connection is lost, the terminal continues to validate '
     'PINs against the cache. Calls between residents and security may be affected, but '
     'gate operations continue. All cached events sync when connectivity is restored.'),
    ('Can residents use their existing phone numbers?',
     'AnfieldVoice uses SIP \u2014 it operates over data (Wi-Fi or cellular), not traditional '
     'phone networks. Residents use the AnfieldVoice app. Their existing phone numbers '
     'are unaffected. The app can also be configured to forward calls to cellular numbers '
     'if desired.'),
    ('How secure are the PINs?',
     'PINs are six-digit, valid for one day only, and hashed in the database (bcrypt). '
     'A PIN can only be used once, for one ticket. Brute-force protection locks the '
     'terminal after repeated invalid attempts. PINs cannot be reverse-engineered from '
     'the database.'),
    ('What about residents who don\'t use smartphones?',
     'For residents without smartphones, we offer alternative workflows: (1) they can '
     'call security to create a visitor ticket on their behalf, (2) they can create tickets '
     'via a web portal, or (3) unticketed visitors can call the resident from the gate '
     'terminal using the apartment number. The system is designed for smartphone-first '
     'operation but does not exclude non-smartphone residents.'),
    ('Can we customise the app with our estate\'s branding?',
     'Yes. Each estate deployment receives a custom-branded build of the AnfieldVoice app '
     'with the estate\'s logo, name, and colour scheme. The app appears under your estate\'s '
     'name in the Apple App Store and Google Play Store.'),
    ('How does this compare to ButterflyMX?',
     'ButterflyMX is a US-focused intercom replacement with basic visitor management. '
     'AnfieldVoice includes everything ButterflyMX offers plus: recurring visitor profiles, '
     'full delivery management, a resident directory, estate broadcasts, custom branding, '
     'on-premises deployment option, multi-language support for South African languages, '
     'and local SA-based support. See the competitive comparison table in Section 15.'),
    ('What does it cost?',
     'Please see the Commercial Proposal in Section 20. Pricing is based on unit count, '
     'deployment model, and selected features. We provide transparent per-unit pricing '
     'with no hidden costs.'),
    ('How long does implementation take?',
     '10\u201314 weeks from contract signing to full operational deployment. See the '
     'Implementation Plan in Section 16 for the detailed phase breakdown.'),
]

for question, answer in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {question}\n')
    run.bold = True
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'A: {answer}')
    run2.font.size = Pt(10.5)

doc.add_page_break()

# ============================================================
# 19. FUTURE ROADMAP
# ============================================================
doc.add_heading('19. Future Roadmap', level=1)

doc.add_paragraph(
    'AnfieldVoice is under active development. The following capabilities are on our '
    '12\u201324 month roadmap. All roadmap items are developed in consultation with our '
    'estate partners to ensure they address real operational needs.'
)

doc.add_heading('19.1 Near-Term (6\u201312 Months)', level=2)

add_colored_table(doc,
    ['Feature', 'Description', 'Status'],
    [
        ['Gate & Boom Integration', 'Automated boom gate triggering on successful PIN validation', 'In Development'],
        ['CCTV Integration', 'Camera snapshot captured on visitor arrival and attached to ticket', 'In Development'],
        ['QR Code Visitor Passes', 'Printable QR codes for events, functions, and long-term contractors', 'Planned'],
        ['SMS PIN Delivery', 'Automated SMS delivery of PINs to visitors from within the app', 'Planned'],
        ['Advanced Analytics', 'Predictive analytics for staffing, peak hour forecasting, anomaly detection', 'Planned'],
    ]
)

doc.add_heading('19.2 Medium-Term (12\u201324 Months)', level=2)

add_colored_table(doc,
    ['Feature', 'Description'],
    [
        ['Number Plate Recognition', 'Automatic vehicle identification on approach \u2014 linked to recurring visitor profiles'],
        ['Biometric Integration', 'Fingerprint or facial recognition option for recurring visitors (resident opt-in)'],
        ['Visitor Self-Service Kiosk', 'Fully self-service visitor check-in for pre-authorised visitors during off-peak hours'],
        ['Multi-Estate Management', 'Centralised management console for property groups with multiple estates'],
        ['Delivery Locker Integration', 'Integration with smart delivery lockers for contactless package handoff'],
    ]
)

doc.add_heading('19.3 Long-Term Vision', level=2)
add_bullet(doc, 'AI-powered anomaly detection \u2014 flagging unusual visitor patterns automatically')
add_bullet(doc, 'Resident community features \u2014 secure neighbourhood social platform')
add_bullet(doc, 'Smart home integration \u2014 intercom, access control, and home automation convergence')
add_bullet(doc, 'Municipal integration \u2014 direct connection to SAPS, armed response, and emergency services')

doc.add_page_break()

# ============================================================
# 20. COMMERCIAL PROPOSAL
# ============================================================
doc.add_heading('20. Commercial Proposal', level=1)

doc.add_paragraph(
    'The following is an indicative commercial proposal for a 360-unit residential estate. '
    'Final pricing is confirmed after the Discovery & Planning phase (Phase 1), when '
    'specific hardware, integration, and customisation requirements are assessed.'
)

doc.add_heading('20.1 Platform Licensing', level=2)

add_colored_table(doc,
    ['Component', 'Model', 'Rate (per unit/month)', 'Annual (360 units)'],
    [
        ['AnfieldVoice Platform (Core)', 'Per-unit SaaS', 'R 45 \u2013 R 65', 'R 194,400 \u2013 R 280,800'],
        ['Resident Mobile App (Branded)', 'Included in platform', '\u2014', '\u2014'],
        ['Security Dashboard', 'Included in platform', '\u2014', '\u2014'],
        ['Administration Portal', 'Included in platform', '\u2014', '\u2014'],
        ['Visitor Terminal License', 'Per terminal', 'R 250 \u2013 R 400', 'R 3,000 \u2013 R 4,800'],
    ]
)

doc.add_paragraph(
    '* Platform licensing is based on total registered units, not active users. '
    'The estate pays a fixed monthly fee regardless of individual resident adoption rates.'
)

doc.add_heading('20.2 One-Time Setup', level=2)

add_colored_table(doc,
    ['Item', 'Indicative Cost', 'Notes'],
    [
        ['Deployment & Configuration', 'R 35,000 \u2013 R 55,000', 'Phases 1\u20133: discovery, infrastructure, platform config'],
        ['Custom App Branding', 'R 15,000 \u2013 R 25,000', 'One-time: custom logo, colours, app store listing setup'],
        ['Staff Training', 'R 8,000 \u2013 R 15,000', 'On-site security staff training (Phase 4)'],
        ['Resident Onboarding Programme', 'R 12,000 \u2013 R 20,000', 'Materials, launch event, support during Phase 5'],
        ['Hardware \u2014 Visitor Terminal', 'R 8,000 \u2013 R 18,000', 'Per terminal. Weatherproof touchscreen + mounting.'],
        ['Hardware \u2014 Security Dashboard PC', 'R 6,000 \u2013 R 12,000', 'Per station. Standard office PC or all-in-one.'],
    ]
)

doc.add_heading('20.3 Ongoing Services', level=2)

add_colored_table(doc,
    ['Service', 'Model', 'Indicative Rate'],
    [
        ['Support & Maintenance (Standard)', 'Per-unit/month', 'R 12 \u2013 R 18'],
        ['Support & Maintenance (Premium)', 'Per-unit/month', 'R 22 \u2013 R 30'],
        ['Cloud Hosting (if applicable)', 'Per-unit/month', 'R 15 \u2013 R 25'],
        ['On-Premises Server Maintenance', 'Monthly retainer', 'R 3,500 \u2013 R 6,000'],
    ]
)

doc.add_heading('20.4 Illustrative Total (360 Units, Cloud-Hosted)', level=2)

add_colored_table(doc,
    ['Category', 'Annual Cost (Range)'],
    [
        ['Platform Licensing', 'R 194,400 \u2013 R 280,800'],
        ['Support & Maintenance (Standard)', 'R 51,840 \u2013 R 77,760'],
        ['Cloud Hosting', 'R 64,800 \u2013 R 108,000'],
        ['Visitor Terminal License (1 gate)', 'R 3,000 \u2013 R 4,800'],
        ['Total Annual Recurring', 'R 314,040 \u2013 R 471,360'],
        ['', ''],
        ['One-Time Setup (incl. hardware)', 'R 84,000 \u2013 R 145,000'],
        ['', ''],
        ['Year 1 Total (all-in)', 'R 398,040 \u2013 R 616,360'],
    ]
)

doc.add_paragraph(
    'Year 1 all-in cost: approximately R 1,100 \u2013 R 1,700 per unit. Year 2 onwards: '
    'approximately R 870 \u2013 R 1,310 per unit per year (recurring only).'
)

doc.add_heading('20.5 Payment Terms', level=2)
add_bullet(doc, 'One-time setup: 50% on contract signing, 50% on go-live')
add_bullet(doc, 'Recurring licensing: monthly or annual invoicing (3% discount for annual prepayment)')
add_bullet(doc, 'Standard payment terms: 30 days from invoice')
add_bullet(doc, 'Minimum commitment: 12 months initial term, then month-to-month')

doc.add_page_break()

# ============================================================
# 21. CONCLUSION
# ============================================================
doc.add_heading('21. Conclusion', level=1)

doc.add_paragraph(
    'AnfieldVoice represents a fundamental shift in how residential estates manage visitors, '
    'deliveries, and resident communication. It is not an intercom upgrade \u2014 it is an '
    'operational platform designed around the way estates actually function.'
)

doc.add_heading('Why AnfieldVoice', level=2)

reasons = [
    ('Purpose-Built',
     'Every feature \u2014 visitor ticketing, recurring visitor profiles, delivery management, '
     'security dashboard \u2014 was designed for residential estates. No generic components '
     'repurposed from office or enterprise systems.'),
    ('Security Without Friction',
     'Multi-layered verification that actually speeds up processing rather than slowing it down. '
     'Residents stay in control. Security stays informed. Visitors move through efficiently.'),
    ('Complete Platform',
     'Voice, video, visitor management, delivery management, resident directory, administration, '
     'and analytics \u2014 one platform, one vendor, one support contract.'),
    ('South African Context',
     'Designed for SA estates: body corporate governance, security officer workflows, multi-language '
     'support, POPIA compliance, local support team, and understanding of the specific visitor '
     'categories and operational patterns of SA residential living.'),
    ('Scalable Foundation',
     'The architecture that serves a 360-unit estate today scales to 5,000+ units without '
     'fundamental rearchitecture. The investment is future-proof.'),
]

for title, desc in reasons:
    p = doc.add_paragraph()
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = TEAL
    p.add_run(desc)

doc.add_paragraph('')
doc.add_paragraph('')

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('We welcome the opportunity to present AnfieldVoice in detail\nand to demonstrate the platform for your estate committee.')
run.font.size = Pt(12)
run.font.color.rgb = DARK
run.italic = True

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Contact\n')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run = p.add_run('[Contact Name]\n[Title]\n[Email Address]\n[Phone Number]\n[Website URL]')
run.font.size = Pt(11)
run.font.color.rgb = DARK

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for word, color in [('SECURE', NAVY), ('  \u25c6  ', RGBColor(0xcc, 0xcc, 0xcc)),
                     ('EFFICIENT', TEAL), ('  \u25c6  ', RGBColor(0xcc, 0xcc, 0xcc)),
                     ('BUILT FOR ESTATES', NAVY)]:
    run = p.add_run(word)
    run.font.size = Pt(12)
    run.font.color.rgb = color
    run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u00a9 2026 AnfieldVoice. All rights reserved. Confidential.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# SAVE
# ============================================================
output_path = '/home/ubuntu/AnfieldVoice_Proposal.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'File size: {os.path.getsize(output_path):,} bytes')
