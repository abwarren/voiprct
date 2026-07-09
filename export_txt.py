from docx import Document
import os

doc = Document('/home/ubuntu/voiprct/AnfieldVoice_Proposal.docx')
lines = []
for p in doc.paragraphs:
    lines.append(p.text)
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append(' | '.join(cells))
    lines.append('')

out = '/home/ubuntu/voiprct/AnfieldVoice_Proposal.txt'
with open(out, 'w') as f:
    f.write('\n'.join(lines))
print(f'Written: {out} ({os.path.getsize(out):,} bytes)')
